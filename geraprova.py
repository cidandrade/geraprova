"""
###########################################################################
geraprova.py
Versão 1.0.3

Autor: Cid R Andrade (profandrade@gmail.com)
Data da versão original: 28/novembro/2023
Data desta versão:       13/dezembro/2023

Como executar:
Em Linux: Abra um terminal, navegue até o diretório deste arquivo e 
execute o seguinte comando
python3 geraprova.py

Caso você não tenha a biblioteca docx, execute o seguinte comando antes 
de usar este programa pela primeira vez
pip install python-docx

Este programa é licenciado como Software Livre: 
Você pode redistribuí-lo e/ou modificá-lo sob os termos da Licença Pública 
Geral GNU, conforme publicada pela Free Software Foundation, seja na versão 
3 da Licença ou (a seu critério) qualquer versão posterior. 
Este programa é distribuído na esperança de que seja útil, mas SEM NENHUMA 
GARANTIA; sem mesmo a garantia implícita de COMERCIABILIDADE ou ADEQUAÇÃO 
A UMA FINALIDADE ESPECÍFICA. 
Veja a Licença Pública Geral GNU para mais detalhes.
Você deveria ter recebido uma cópia da Licença Pública Geral GNU junto com 
este programa. Se não recebeu, consulte https://www.gnu.org/licenses/.

ChangeLog
1.0.0 28/11/2023: Versão operacional
1.0.1 2/12/2023: Otimização de desempenho
1.0.2 5/12/2023: Inclusão de tratamento de erro em abertura de arquivo
1.0.3 13/12/2023: Gera a prova em DOCX, organiza nomes de variáveis, 
                  seleciona arquivo sem precisar mexer no código

Cid R. Andrade
 
Gera questões de múltipla escolha

O arquivo que é importado deve conter uma lista com diversas tuplas de 
questões. Cada tupla deve conter o enunciado e as cinco opções, sendo a 
primeira delas a correta
Exemplo: [("Enunciado", "Correta", "Distrator 1", "Distrator 2", 
           "Distrator 3", "Distrator 4")]
###########################################################################
"""

# A biblioteca ast é utilizada para avaliar se o arquivo externo está 
# corretamente formatado como uma expressão Python válida
import ast
# A biblioteca datetime irá lidar com questões de datas
from datetime import datetime
# A biblioteca docx é utilizada para gerar arquivos DOCX
from docx import Document
# A biblioteca os é para acessar dados do sistema operacional
import os
# A biblioteca random é utilizada para embaralhar questões e opções
import random

# Utilidades
ABC = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
# Nomes das Avaliações
TP_AV = ("Avaliação Intermediária", 
	      "Avaliação Regimental", 
	      "Avaliação Final")
CD_AV = ("AI", "AR", "AF")

# Funções
# A função intInput recebe uma String e força uma entrada numérica 
# inteira válida
def intInput(p):
	while True:
		try:
			v = int(input(p))
			return v
		except ValueError:
			print("Digite um número inteiro válido")

# A função exibirMenu exibe um menu com opções
def exibirMenu(o, s):
	print (s)
	for i, st in enumerate(o, start=1):
		print (f"{i}. {st}")

# A função obterEscolha exibe um menu e obtém uma opção válida
def obterEscolha(o, s):
	q = len(o)
	exibirMenu(o, s)
	while True:
		e = intInput(f"Digite o número da opção desejada (1 a {q}): ")
		if 1 <= e <= q:
			return e
		else:
			print (f"Digite um número válido entre 1 e {q}")

# A função listaArquivos lista arquivos que podem ter questões de provas
def listaArquivos():
	currDir = os.path.dirname(os.path.abspath(__file__))
	return [arq for arq in os.listdir(currDir) if arq.endswith('.txt')]

# Código principal

# Tenta abrir o arquivo com as questões

# Defina o arquivo com as questoes
arquivos = listaArquivos()
selArquivo = obterEscolha(arquivos, 
	"Qual arquivo contém as questões de prova?")
nomeFonte = arquivos[selArquivo - 1]

try:
	with open(nomeFonte, 'r') as arqFonte:
		questoes = ast.literal_eval(arqFonte.read())

  # Pergunta quantos modelos de prova devem ser feitos
	qtProvas = intInput("Deseja quantas provas: ")
	# Prossegue se a quantidade de provas for entre 1 e 26
	if 0 < qtProvas <= len(ABC):
		# Pergunta a quantidade de questões e valor de cada uma
		qtQuest = intInput("Deseja quantas questões por prova: ")
		valQuest = input("Quanto vale cada questão: ")
		# Obtém a quantidade de questões disponíveis
		tamQuestoes = len(questoes)

		# Prossegue se a quantidade de questões desejada for menor ou igual a 
		# quantidade disponível
		if qtQuest <= tamQuestoes:
			# A variável nmLocal conterá o nome de cada uma das provas
			# nmsProvas uma lista com o nome de todas as provas
			nmLocal, nmsProvas = "", []
			# Obtém tipo da avaliação
			nTipo = obterEscolha(TP_AV, "Qual é o tipo da prova?")
			tipo = TP_AV[nTipo - 1]
			# Obtém semestre e ano atuais
			semestre = (datetime.now().month - 1) // 6 + 1
			ano = datetime.now().year
			# Obtém o nome da disciplina
			disciplina = input("Qual é o nome da disciplina? ")

			# Cria as provas com nomes A, B, C e assim por diante
			for nome in ABC[:qtProvas]:
				# A variável prova conterá o texto da prova
				prova = ""
				# Obtém lista com números únicos das questões selecionadas
				qSelec = random.sample(range(tamQuestoes-1), qtQuest)
				# Nome desta prova
				nmLocal = f"Prova_{CD_AV[nTipo - 1]}_{nome}.docx"
				# Adiciona na lista com nomes das provas
				nmsProvas.append(nmLocal)
				# Abre o modelo de prova
				docmnt = Document("ModeloProvaUnicid.docx")
				# Obtém parágrafos do modelo
				para = list(docmnt.paragraphs)
				# Escreve nome da disciplina
				para[3].text = f"Disciplina: {disciplina}"
				# Escreve a identificação da prova (Ex.: 'Prova A') e a palavra 
				# 'Gabarito'
				prova += f"Prova {nome}\nGabarito\n"
				# A variável qProva conterá as questões da prova
				qProva = ""

				# Laço para tratar cada uma das questões
				for qtQuestao in range(qtQuest):
					# Escreve o número da questão e seu valor
					# Exemplo: 'QUESTÃO 1 - 0,25 PONTO'
					qProva += f"\nQUESTÃO {qtQuestao + 1} - {valQuest} PONTO\n\n"
					# Escreve o enunciado da questão
					qProva += questoes[qSelec[qtQuestao]][0] + '\n\n'
					# Obtém uma lista com a ordem das opções
					ordemOpcoes = random.sample(range(1,6),5)
					# A variável 'opcoes' terá as opções da questão
					opcoes = ""

					# Laço para lidar com as opções da questão
					for iOpc, j in enumerate(ordemOpcoes):
						# Escreve cada opção
						# Exemplo: '(A) blá, blá, blá'
						opcoes += f"({ABC[iOpc]}) {questoes[qSelec[qtQuestao]][j]}\n"

						# Se a opção for a primeira delas (com índice '1'), indicando 
						# ser a correta
						if (ordemOpcoes[j-1] == 1):
							# Escreva a entrada do gabarito da questão
							# Exemplo: '1: A'
							prova += f"{qtQuestao + 1}: {ABC[j-1]}\n"

					# Appenda na questão as opções dela
					qProva += opcoes

				# Appenda na prova todas as questões
				prova += "Questões\n" + qProva
				
				# Escreve arquivo com o conteúdo da prova
				docmnt.add_paragraph(prova)
				docmnt.save(nmLocal)

			# Informa quais arquivos foram gerados
			print ("Provas armazenadas nos arquivos")
			print (*nmsProvas, sep = ", ")

		else:
			print(f"Tenho somente {tamQuestoes} questões disponíveis")

	else:
		print (f"Não é possível gerar {qtProvas} provas")

except FileNotFoundError:
	print (f"O arquivo {nomeFonte} não foi encontrado!!!")
except Exception as e:
	print (f"Ocorreu um erro: {e}")